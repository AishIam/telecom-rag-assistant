"""
Offline ingestion: load documents, chunk them, embed, and store in ChromaDB.
Run once (or whenever source documents change).
"""

import re
from pathlib import Path
from docx import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma

import config


def iter_block_items(parent):
    """Yield paragraphs and tables in document order."""
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def extract_tables_from_docx(path):
    """Extract DOCX table rows as individual Documents."""
    doc = DocxDocument(str(path))
    table_docs = []

    current_section = ""
    current_table_caption = ""
    table_index = 0

    for block in iter_block_items(doc):

        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue

            if text.startswith("Table "):
                current_table_caption = text
                continue

            style_name = block.style.name.lower() if block.style else ""

            if "heading" in style_name:
                current_section = text
            elif re.match(r"^(?:\d+(?:\.\d+)*|[A-Z](?:\.\d+)*)", text):
                current_section = text

            continue

        if isinstance(block, Table):

            rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]

            if len(rows) < 2:
                table_index += 1
                continue

            if len(rows) > 2:
                num_cols = max(len(rows[0]), len(rows[1]))
                headers = []

                for col in range(num_cols):
                    parts = []
                    for hr in rows[:2]:
                        if col < len(hr):
                            value = hr[col].strip()
                            if value and value not in parts:
                                parts.append(value)
                    headers.append(" | ".join(parts))

                data_rows = rows[2:]
            else:
                headers = rows[0]
                data_rows = rows[1:]

            for row_index, row in enumerate(data_rows):

                if sum(bool(c.strip()) for c in row) < 2:
                    continue

                if " ".join(row).lower().startswith("note"):
                    continue

                pairs = []
                for h, v in zip(headers, row):
                    h = h.strip()
                    v = v.strip()
                    if h and v:
                        pairs.append(f"{h}: {v}")

                if not pairs:
                    continue

                pairs = list(dict.fromkeys(pairs))
                row_text = "\n".join(pairs)

                page_text = f"""3GPP Specification

Source File:
{path.name}

Section:
{current_section}

Table:
{current_table_caption}

Table Index:
{table_index}

{row_text}"""

                table_docs.append(
                    Document(
                        page_content=page_text,
                        metadata={
                            "source_file": path.name,
                            "content_type": "table",
                            "document_type": "3gpp_table",
                            "table_index": table_index,
                            "row_index": row_index,
                            "section": current_section,
                            "table_caption": current_table_caption,
                        },
                    )
                )

            table_index += 1

    return table_docs


def load_documents(raw_dir):
    docs = []

    for path in raw_dir.iterdir():

        if path.suffix.lower() == ".pdf":
            loaded = PyPDFLoader(str(path)).load()

            for d in loaded:
                d.metadata["source_file"] = path.name
                d.metadata["content_type"] = "text"

            docs.extend(loaded)
            print(f"[LOAD] {path.name} -> {len(loaded)} page(s)")

        elif path.suffix.lower() == ".docx":
            doc = DocxDocument(str(path))

            paragraphs = [
                p.text.strip()
                for p in doc.paragraphs
                if p.text.strip()
            ]

            if paragraphs:
                docs.append(
                    Document(
                        page_content="\n\n".join(paragraphs),
                        metadata={
                            "source_file": path.name,
                            "content_type": "text",
                        },
                    )
                )

            table_docs = extract_tables_from_docx(path)
            docs.extend(table_docs)

            print(f"[LOAD] {path.name} -> 1 text document(s) + {len(table_docs)} table row(s)")

        else:
            print(f"[SKIP] {path.name}")

    return docs


def chunk_documents(docs):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    chunks = []

    for doc in docs:
        if doc.metadata.get("content_type") == "table":
            chunks.append(doc)
        else:
            chunks.extend(splitter.split_documents([doc]))

    print(f"[CHUNK] {len(docs)} document(s) -> {len(chunks)} chunk(s)")
    return chunks


def build_vectorstore(chunks):
    embeddings = HuggingFaceEmbeddings(model_name=config.EMBEDDING_MODEL)

    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=config.COLLECTION_NAME,
        persist_directory=str(config.VECTOR_DB_DIR),
    )

    print(f"[STORE] Persisted {len(chunks)} chunks to {config.VECTOR_DB_DIR}")
    return vectordb


if __name__ == "__main__":
    docs = load_documents(config.RAW_DATA_DIR)

    if not docs:
        raise SystemExit("No documents found in data/raw/")

    chunks = chunk_documents(docs)
    build_vectorstore(chunks)

    print("[DONE] Ingestion Complete!")